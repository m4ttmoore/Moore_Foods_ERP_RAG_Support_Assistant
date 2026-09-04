# load_document.py
#
# Purpose: extract plain text from a Word (.docx) source document, including
# both paragraphs AND table content, so it can be chunked and embedded later.
#
# Why tables matter: python-docx's default doc.paragraphs only gives you
# paragraph text. If your source document has any tables (status definitions,
# role lists, anything in a grid), that content is invisible unless you walk
# the document body manually and check for both paragraph and table elements,
# which is what this script does below.

from docx import Document

def load_docx_text(path):
    doc = Document(path)
    parts = []
    # doc.element.body gives us elements in their ORIGINAL document order,
    # which matters: it keeps a table's content positioned correctly relative
    # to the surrounding paragraphs, rather than dumping all paragraphs first
    # and all tables afterwards.
    for element in doc.element.body:
        if element.tag.endswith('}p'):
            # paragraph
            for para in doc.paragraphs:
                if para._element is element:
                    if para.text.strip():
                        parts.append(para.text)
                    break
        elif element.tag.endswith('}tbl'):
            for table in doc.tables:
                if table._element is element:
                    for row in table.rows:
                        # Flattens each table row into one line, cells joined
                        # with " | ". Adapt this if your own tables need a
                        # different flattened format for the chunker/LLM to
                        # read cleanly, e.g. more descriptive labels than a
                        # bare pipe-separated row.
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        parts.append(row_text)
                    break
    return "\n".join(parts)

if __name__ == "__main__":
    # PROJECT SPECIFIC: this hardcoded path is the one thing you MUST change
    # for a different project. Point it at your own source .docx file.
    # If you have multiple source documents, this is also the place you'd
    # extend into a loop over a folder rather than one fixed filename.
    text = load_docx_text("Data/Moore_Foods_ERP_Process_Knowledge_Base_v3.docx")

    # encoding="utf-8" here (and everywhere this project writes/reads text)
    # is not optional on Windows. Without it, Python defaults to cp1252,
    # which crashes on non-ASCII characters like arrows (→) or curly quotes
    # that a Word document commonly contains. If you hit a
    # UnicodeEncodeError/UnicodeDecodeError anywhere else in a project like
    # this, this is almost always the fix.
    with open("knowledge_base.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print(f"Extracted {len(text)} characters.")
